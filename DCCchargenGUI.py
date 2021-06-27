# -*- coding: utf-8 -*-
"""
last updated 10/06/2021

@author: P Beeltje
"""

from openpyxl import load_workbook
import random
import PySimpleGUI as sg
from docx import Document
from datetime import datetime

def charGen(characters, sets):
    
    
    #load excel file and specific sheets
    wb = load_workbook('DCCsheet.xlsx')
    EquipSheet = wb['equipment']
    OccuSheet = wb['occupation']
    QuirkSheet = wb['quirks2']
    AugurSheet = wb['signs']
    
    document = Document() 
    
    #this function sets count to the nr of filled in cells in column A of this sheet
    def rowCounter(sheet):
        count = 0
        x = 1
        while True:
            cell = 'A'+ str(x)
            if sheet[cell].value is not None:
                count += 1
                x+=1
            else: break
        return count
    
    #python switch type that picks modifier based on stat
    def modget(mod):
            switcher={
                    3:-3, 4:-2, 5:-2, 6:-1, 7:-1, 8:-1, 9:0, 10:0, 11:0, 12:1, 13:1, 14:1, 15:1, 16:2, 17:2, 18:3
                 }
            return switcher.get(mod,"Invalid AbilityScore")
    
    # rolls and adds 3 6-sided die
    def abilityroll ():
        roll = random.randint(1, 6) + random.randint(1, 6) + random.randint(1, 6)
        return roll
    
    def abilitybonusroll ():
        roll = random.randint(1, 6) + random.randint(1, 6) + 6
        return roll
    
    def namegen ():
            part1 = ('Kei', 'Ber', 'Yu', 'Con', 'Kaji', 'Sam','Fro','Gan', 'Bun', 'Dai', 'Yuna', 'Sar', 'Shige', 'Phil', 'Geirr', 'Fride', 'On', 'Ike', 'Xan', 'Æthel', 'Alf', 'Wil', 'Ma', 'Chi', 'Yo', 'Ba', 'Hu', 'Mitsu', 'Ai', 'Al', 'Pen', 'Æðel')
            part2 = ('ko', 'ki', 'ka', 'zumi', 'zaki', 'ger','agald', 'kiko', 'hippus', 'ram','wise', 'do', 'vasius', 'nin', 'ji', 'swin', 'wald', 'malric', 'ran', 'stan', 'ron', 'dalf', 'stein', 'michi', 'dert', 'bert',  'hiki', '', 'shi',  'mochi', 'trand', 'rod', 'eger', 'agon', 'son','eron', 'lip', 'do', 'cras', 'red', 'hiko', 'koto')
            Namepart1 = random.choice(part1)
            Namepart2 = random.choice(part2)
            Firstname = Namepart1 + Namepart2
            last1 = ('the ', 'son of ', '', '', '')
            last3 = ('Yokohama', 'Michitsune', 'Naotomo', 'Wada', 'Ide', 'Higasa', 'Suzuki', 'Sato', 'Ito', 'Teruhira','Kagekazu', 'Shigetoki', 'Munetami', 'Bakemono', 'Sukeyasu', 'Taiko', 'Akitoki', 'Yamada Bome', 'Cazellnu', 'Kujo', 'Yoshi', 'Murai', 'Kircheis','von Reuenthal', 'von Mariendorf', 'Mittermeyer', 'von Oberstein', 'Bittenfeld', 'Kesler', 'Yang', 'von Fahrenheit', 'von Eisenach')
            thelast = ('Bold', 'Bald', 'Branded', 'Wise', 'Quick', 'Kind', 'Silent', 'Sizeless', 'Peerless', 'Stranger', 'Watcher', 'Unwise', 'Numbered', 'Believer', 'Bastard', 'Tall', 'Cursed', 'Lucky', 'Unlucky', 'Blessed')
            
            Lastpart1 = random.choice(last1)
            if Lastpart1 == 'the ':
                lastpartthe = random.choice(thelast)
                Lastname = 'the ' + lastpartthe
            else:
                Lastpart3 = random.choice(last3)
                Lastname = Lastpart1 + Lastpart3
            Name = Firstname + " " + Lastname
            return Name
    
    
    equipcount = rowCounter(EquipSheet) #counts equipment options
    occucount = rowCounter(OccuSheet) #counts occupation options
    quirkcount = rowCounter(QuirkSheet)
    augurcount = rowCounter(AugurSheet)
    
    #Starting questions pick a number or characters and sets of characters:
    chartotal = int(characters)
    settotal = int(sets)
    setascending = 1
    #Loops HERE
    
    while settotal > 0 :
        print("Set ", setascending, '\n')
        document.add_heading('Set ' + str(setascending), 1)
        chardescending = chartotal
        charascending = 1 #had to create ascending variables to be able to print char and set numbers!
        while chardescending > 0:
    
            #ability+modifiers
            
            #bonus = certain column in count row on occupation sheet
            #and then if bonus = S strength = abilitybonuxroll()
            Name = namegen()
            occurow = random.randint(1, occucount)
            quirkrow = random.randint(1, quirkcount)
            augurrow = random.randint(1, augurcount)
    
            Occupation = OccuSheet[('A' + str(occurow))].value
            Quirk = QuirkSheet[('A' + str(quirkrow))].value
            Augur = AugurSheet['A' + str(augurrow)].value
            
            if OccuSheet[('D' + str(occurow))].value == 'Strength':
               Strength = abilitybonusroll() 
            else:
                Strength = abilityroll()
            StrengthMod = modget(Strength)
            
            if OccuSheet[('D' + str(occurow))].value == 'Agility':
               Agility = abilitybonusroll() 
            else:
                Agility = abilityroll()
            AgilityMod = modget(Agility)
            
            if OccuSheet[('D' + str(occurow))].value == 'Stamina':
               Stamina = abilitybonusroll() 
            else:
                Stamina = abilityroll()
            StaminaMod = modget(Stamina)
            if OccuSheet[('D' + str(occurow))].value == 'Personality':
               Personality = abilitybonusroll() 
            else:
                Personality = abilityroll()
            PersonalityMod = modget(Personality)
            if Personality < 10:
                Personality = '  ' + str(Personality) #for formatting reasons
            
            if OccuSheet[('D' + str(occurow))].value == 'Intelligence':
               Intelligence = abilitybonusroll() 
            else:
                Intelligence = abilityroll()
            IntelligenceMod = modget(Intelligence)
            
            if OccuSheet[('D' + str(occurow))].value == 'Luck':
               Luck = abilitybonusroll() 
            else:
                Luck = abilityroll()
            LuckMod = modget(Luck)
            
            #other stats
            HP= random.randint(1,4) + StaminaMod
            if HP < 1:
                continue
            AC= 10 + AgilityMod
            Weapon = OccuSheet[('B' + str(occurow))].value
            Tradegood = OccuSheet[('C' + str(occurow))].value
            Equipmentno = 'A' + str(random.randrange(1, equipcount)) #generates a cell-code column+random number
            Equipment = EquipSheet[Equipmentno].value
            Wealth= random.randint(1, 12) + random.randint(1, 12) + random.randint(1, 12) + random.randint(1, 12) + random.randint(1, 12)
            
            #Saves are same as 3 ability mods
            
            #printing!
            print(charascending, end=' ')
            print(Name)
            print("Occupation:", Occupation, '\n')
            print("HP:", HP)
            print ("Unarmoured AC:", AC)
            print("Strength:    ", Strength, " Mod: ", StrengthMod)
            print("Agility:     ", Agility, " Mod: ", AgilityMod)
            print("Stamina:     ", Stamina, " Mod: ", StaminaMod)
            print("Personality: ", Personality, " Mod: ", PersonalityMod)
            print("Intelligence:", Intelligence, " Mod: ", IntelligenceMod)
            print("Luck:        ", Luck, " Mod: ", LuckMod, '\n') 
            print("Saves:") 
            print("Fortitude:", StaminaMod, "Reflex:", AgilityMod, "Will:", PersonalityMod, '\n')
            print("Init:", AgilityMod, '\n')
            print("Equipment:\n", Weapon, '\n', Equipment, '\n', Tradegood, '\n',  Wealth, "copper pieces", '\n\n')
            print("Quirk: \n", Quirk, "\n \n")
            print("Birth Augur: ", Augur)
            
            #document generation
    
            nl= '\n'
            tab= '\t'
            document.add_heading(str(charascending) + ' ' + Name + ', The ' + Occupation, 2)
            Para = (
      f"{nl}Strength:     {Strength} Mod: {StrengthMod} {tab}{tab}HP: {HP}{nl}"
      f"Agility:      {Agility} Mod: {AgilityMod}{tab}{tab}AC: {AC}{nl}"
      f"Stamina:      {Stamina} Mod: {StaminaMod}{nl}"
      f"Personality:  {Personality} Mod: {PersonalityMod}{tab}SAVES:{nl}"
      f"Intelligence: {Intelligence} Mod: {IntelligenceMod}{tab}{tab}Fortitude: {StaminaMod}  Reflex: {AgilityMod}  Will: {PersonalityMod}{nl}"
      f"Luck:         {Luck} Mod: {LuckMod}{tab}{tab}Init bonus: {AgilityMod}{nl}{nl}"
      f"EQUIPMENT{nl}{Weapon}{nl}{Equipment}{nl}{Tradegood}{nl}{Wealth} copper pieces{nl}{nl}"
      f"Birth Augur(luck bonus): {Augur}{nl}{nl}"
      f"Quirk: {Quirk}{nl}{nl}"
                  )
            paragraph = document.add_paragraph(Para)
            #paragraph.style = 'Subtitle'
    
            
            
            chardescending -= 1
            charascending += 1
        settotal -= 1
        setascending += 1
    
    print("\n END \n ")
    dateTimeObj = datetime.now() #timestamp to make a unique document name
    timestampStr = dateTimeObj.strftime("%d %b  %H.%M.%S")
    docname = 'DCCgen ' + timestampStr +'.docx' 
    document.save(docname)
    print("Generated document")
    return docname




#GUI


characters = 4
sets = 2

layout = [  [sg.Text('Generate a docx file with random characters')],
            [sg.Text('Enter Number of Characters per Set:')],
            [sg.InputText(key= "characters", default_text="4")],
            [sg.Text('Enter Number of Sets:')],
            [sg.InputText(key= "sets", default_text="2")],
            [sg.Text('', key='console', size=(40,1))],
            [sg.Button("Generate")]]

window = sg.Window("DCCgen", layout)

while True:
       event, values = window.read()
       if event == sg.WIN_CLOSED:
           break
       if event == "Generate":
           try:
              characters= int(values["characters"])
              sets = int(values["sets"])
           except ValueError: 
              message = "error: please enter a number"
              window['console'].update(message)
           
           message = "Generated " + str(charGen(characters, sets))
           window['console'].update(message)


window.close()


